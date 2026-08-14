import io

from tests.helpers import wait_for_document_status


def _build_minimal_pdf_bytes(text: str) -> bytes:
    """Erzeugt ein minimales, aber korrekt referenziertes Ein-Seiten-PDF mit
    extrahierbarem Text (echte xref-Tabelle mit byte-genauen Offsets) - fuer
    Tests ohne zusaetzliche Abhaengigkeit wie reportlab.
    """
    content_stream = f"BT /F1 12 Tf 20 100 Td ({text}) Tj ET".encode("latin-1")
    objects = [
        b"<< /Type /Catalog /Pages 2 0 R >>",
        b"<< /Type /Pages /Kids [3 0 R] /Count 1 >>",
        b"<< /Type /Page /Parent 2 0 R /Resources << /Font << /F1 4 0 R >> >> "
        b"/MediaBox [0 0 200 200] /Contents 5 0 R >>",
        b"<< /Type /Font /Subtype /Type1 /BaseFont /Helvetica >>",
        b"<< /Length " + str(len(content_stream)).encode() + b" >>\nstream\n"
        + content_stream
        + b"\nendstream",
    ]

    body = bytearray(b"%PDF-1.4\n")
    offsets = [0]
    for i, obj in enumerate(objects, start=1):
        offsets.append(len(body))
        body += f"{i} 0 obj".encode() + b"\n" + obj + b"\nendobj\n"

    xref_offset = len(body)
    xref = f"xref\n0 {len(objects) + 1}\n".encode() + b"0000000000 65535 f \n"
    for off in offsets[1:]:
        xref += f"{off:010d} 00000 n \n".encode()

    trailer = (
        f"trailer\n<< /Size {len(objects) + 1} /Root 1 0 R >>\n"
        f"startxref\n{xref_offset}\n%%EOF"
    ).encode()

    return bytes(body) + xref + trailer


_MINIMAL_PDF_BYTES = _build_minimal_pdf_bytes("Ansprechpartner ist Frau Weber.")


def _build_docx_bytes(paragraphs: list[tuple[str, int | None]]) -> bytes:
    """paragraphs: Liste von (Text, Ueberschriftenebene) - Ebene None = normaler Absatz."""
    import docx

    document = docx.Document()
    for text, heading_level in paragraphs:
        if heading_level:
            document.add_heading(text, level=heading_level)
        else:
            document.add_paragraph(text)
    buf = io.BytesIO()
    document.save(buf)
    return buf.getvalue()


def _create_project(client) -> int:
    return client.post("/api/projects", json={"name": "Upload-Testprojekt"}).json()["id"]


def test_upload_txt_document_becomes_ready(client):
    project_id = _create_project(client)
    content = "Der SAP-Support-Ansprechpartner fuer VA02 ist Frau Weber.".encode("utf-8")

    response = client.post(
        f"/api/projects/{project_id}/documents/upload",
        files={"file": ("notiz.txt", content, "text/plain")},
        data={"typ": "notiz"},
    )

    assert response.status_code == 201
    created = response.json()
    assert created["status"] == "pending"  # Extraktion+Indexierung laufen asynchron (Schritt 8)
    assert created["dateiname"] == "notiz.txt"
    assert created["titel"] == "notiz"  # aus Dateinamen abgeleitet, da kein Titel angegeben

    body = wait_for_document_status(client, project_id, created["id"])
    assert body["status"] == "ready"
    assert "Frau Weber" in body["inhalt"]
    assert body["dokumentdatum"] is not None  # Upload-Zeitpunkt als Vorbelegung


def test_upload_pdf_document_extracts_text(client):
    project_id = _create_project(client)

    response = client.post(
        f"/api/projects/{project_id}/documents/upload",
        files={"file": ("bericht.pdf", _MINIMAL_PDF_BYTES, "application/pdf")},
        data={"typ": "notiz", "titel": "Bericht"},
    )

    assert response.status_code == 201
    created = response.json()

    body = wait_for_document_status(client, project_id, created["id"])
    assert body["status"] == "ready"
    assert "Weber" in body["inhalt"]


def test_upload_docx_document_preserves_headings_for_chunking(client, app):
    project_id = _create_project(client)
    docx_bytes = _build_docx_bytes(
        [
            ("Abschnitt Eins", 1),
            ("Hier steht der Text zum ersten Abschnitt. " * 10, None),
            ("Abschnitt Zwei", 1),
            ("Hier steht der Text zum zweiten Abschnitt. " * 10, None),
        ]
    )

    response = client.post(
        f"/api/projects/{project_id}/documents/upload",
        files={
            "file": (
                "prozess.docx",
                docx_bytes,
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )
        },
        data={"typ": "prozess"},
    )

    assert response.status_code == 201
    created = response.json()

    body = wait_for_document_status(client, project_id, created["id"])
    assert body["status"] == "ready"
    assert "# Abschnitt Eins" in body["inhalt"]
    assert "# Abschnitt Zwei" in body["inhalt"]

    from app.db.models.chunk import Chunk

    with app.state.session_factory() as db:
        chunks = db.query(Chunk).filter(Chunk.document_id == body["id"]).all()
    # Struktur-Chunker sollte anhand der Ueberschriften mehrere Chunks erzeugt haben.
    assert len(chunks) >= 2
    assert any(c.abschnitt == "Abschnitt Eins" for c in chunks)


def test_upload_rejects_unsupported_extension(client):
    project_id = _create_project(client)
    response = client.post(
        f"/api/projects/{project_id}/documents/upload",
        files={"file": ("bild.png", b"\x89PNG\r\n\x1a\n", "image/png")},
        data={"typ": "notiz"},
    )
    assert response.status_code == 422


def test_upload_rejects_content_mismatching_extension(client):
    project_id = _create_project(client)
    response = client.post(
        f"/api/projects/{project_id}/documents/upload",
        files={"file": ("faelschung.pdf", b"Das ist gar kein PDF.", "application/pdf")},
        data={"typ": "notiz"},
    )
    assert response.status_code == 422


def test_upload_enforces_max_size(client, monkeypatch):
    from app.config import get_settings

    monkeypatch.setenv("MAX_UPLOAD_SIZE_MB", "0")  # 0 MB -> jede Datei ist zu gross
    get_settings.cache_clear()

    project_id = _create_project(client)
    response = client.post(
        f"/api/projects/{project_id}/documents/upload",
        files={"file": ("notiz.txt", b"Ein kurzer Text.", "text/plain")},
        data={"typ": "notiz"},
    )
    assert response.status_code == 422
    get_settings.cache_clear()


def test_duplicate_upload_is_stopped_unless_forced(client):
    project_id = _create_project(client)
    content = b"Immer derselbe Inhalt."

    first = client.post(
        f"/api/projects/{project_id}/documents/upload",
        files={"file": ("original.txt", content, "text/plain")},
        data={"typ": "notiz"},
    )
    assert first.status_code == 201

    duplicate = client.post(
        f"/api/projects/{project_id}/documents/upload",
        files={"file": ("kopie.txt", content, "text/plain")},
        data={"typ": "notiz"},
    )
    assert duplicate.status_code == 409

    forced = client.post(
        f"/api/projects/{project_id}/documents/upload",
        files={"file": ("kopie.txt", content, "text/plain")},
        data={"typ": "notiz", "force": "true"},
    )
    assert forced.status_code == 201
    assert forced.json()["id"] != first.json()["id"]


def test_download_uploaded_file_returns_original_content(client):
    project_id = _create_project(client)
    content = b"Originalinhalt zum Herunterladen."

    created = client.post(
        f"/api/projects/{project_id}/documents/upload",
        files={"file": ("download-test.txt", content, "text/plain")},
        data={"typ": "notiz"},
    ).json()

    response = client.get(f"/api/projects/{project_id}/documents/{created['id']}/file")
    assert response.status_code == 200
    assert response.content == content
    assert "download-test.txt" in response.headers["content-disposition"]


def test_upload_for_unknown_project_returns_404(client):
    response = client.post(
        "/api/projects/999999/documents/upload",
        files={"file": ("notiz.txt", b"Text", "text/plain")},
        data={"typ": "notiz"},
    )
    assert response.status_code == 404
